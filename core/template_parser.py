"""
Извлечение переменных Jinja2 из Word-шаблона .docx через AST.
Без выполнения шаблона; только парсинг и список имён переменных.
"""
from dataclasses import dataclass, field as dc_field
from pathlib import Path
from typing import List, Set, Tuple

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from jinja2 import Environment, meta
import jinja2.nodes as j2nodes

from logger import py_logger


@dataclass
class LoopBlock:
    loop_var: str
    item_var: str
    fields: List[str]


@dataclass
class TemplateStructure:
    simple_fields: List[str]
    loop_blocks: List[LoopBlock]
    errors: List[str] = dc_field(default_factory=list)


def _collect_text_from_docx(doc: Document) -> str:
    """Собрать весь текст из параграфов и ячеек таблиц в один строковый шаблон."""
    parts: List[str] = []
    for block in doc.element.body:
        if block.tag.endswith("p"):
            # Один параграф
            p = Paragraph(block, doc)
            parts.append(p.text)
        elif block.tag.endswith("tbl"):
            # Таблица
            tbl = Table(block, doc)
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        parts.append(p.text)
    return "\n".join(parts)


def extract_variables_from_string(template_string: str) -> Tuple[Set[str], List[str]]:
    """
    Извлечь имена переменных из строки шаблона через Jinja2 AST.
    Возвращает (множество имён переменных, список ошибок валидации).
    Для MVP допускаются только простые подстановки {{ name }}.
    """
    env = Environment()
    errors: List[str] = []
    try:
        ast = env.parse(template_string)
    except Exception as e:
        errors.append(f"Синтаксическая ошибка Jinja2: {e}")
        return set(), errors

    variables = meta.find_undeclared_variables(ast)
    # Нормализуем точечную нотацию в плоский ключ для контекста (docxtpl принимает вложенные dict)
    return variables, errors


def extract_variables_from_docx(path_or_stream) -> Tuple[List[str], List[str]]:
    """
    Загрузить .docx, собрать текст, извлечь переменные через Jinja2 AST.
    path_or_stream: путь к файлу (str | Path) или file-like объект.
    Возвращает (список имён переменных, список ошибок).
    """
    errors: List[str] = []
    try:
        if hasattr(path_or_stream, "read"):
            doc = Document(path_or_stream)
        else:
            doc = Document(str(path_or_stream))
    except Exception as e:
        errors.append(f"Не удалось открыть документ: {e}")
        return [], errors

    try:
        template_text = _collect_text_from_docx(doc)
    except Exception as e:
        errors.append(f"Ошибка чтения текста из docx: {e}")
        return [], errors

    variables, parse_errors = extract_variables_from_string(template_text)
    errors.extend(parse_errors)
    # Сортируем для стабильного порядка
    return sorted(variables), errors


def validate_template_syntax(path_or_stream) -> Tuple[bool, List[str]]:
    """
    Проверить, что шаблон допустим для MVP (только {{ var }}).
    Возвращает (ok, список сообщений об ошибках).
    """
    names, errors = extract_variables_from_docx(path_or_stream)
    return len(errors) == 0, errors


def _iter_ast_nodes(node):
    yield node
    for child in node.iter_child_nodes():
        yield from _iter_ast_nodes(child)


def _collect_loop_fields(body_nodes: list, item_var: str) -> List[str]:
    seen: dict = {}
    for top in body_nodes:
        for node in _iter_ast_nodes(top):
            if (
                isinstance(node, j2nodes.Getattr)
                and isinstance(node.node, j2nodes.Name)
                and node.node.name == item_var
            ):
                seen[node.attr] = True
    return list(seen.keys())


def extract_template_structure(path_or_stream) -> "TemplateStructure":
    errors: List[str] = []
    try:
        if hasattr(path_or_stream, "read"):
            doc = Document(path_or_stream)
        else:
            doc = Document(str(path_or_stream))
        template_text = _collect_text_from_docx(doc)
    except Exception as e:
        return TemplateStructure([], [], [str(e)])

    env = Environment()
    try:
        ast = env.parse(template_text)
    except Exception as e:
        return TemplateStructure([], [], [f"Синтаксическая ошибка Jinja2: {e}"])

    loop_blocks: List[LoopBlock] = []
    loop_iter_vars: Set[str] = set()

    for node in ast.body:
        if not isinstance(node, j2nodes.For):
            continue
        if not isinstance(node.iter, j2nodes.Name):
            errors.append(
                f"Цикл со сложным итератором не поддерживается "
                f"(строка {getattr(node, 'lineno', '?')}), блок пропущен."
            )
            continue
        if not isinstance(node.target, j2nodes.Name):
            errors.append(
                f"Цикл с деструктуризацией не поддерживается "
                f"(строка {getattr(node, 'lineno', '?')}), блок пропущен."
            )
            continue
        loop_var = node.iter.name
        item_var = node.target.name
        fields = _collect_loop_fields(node.body, item_var)
        loop_blocks.append(LoopBlock(loop_var=loop_var, item_var=item_var, fields=fields))
        loop_iter_vars.add(loop_var)

    all_vars = meta.find_undeclared_variables(ast)
    simple_fields = sorted(all_vars - loop_iter_vars)

    return TemplateStructure(simple_fields=simple_fields, loop_blocks=loop_blocks, errors=errors)
