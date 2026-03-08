"""
Создание тестового Word-шаблона .docx с плейсхолдерами Jinja2.
Поля совпадают с тестовыми данными из init_db (Договор оказания услуг).

Запуск из корня проекта:
  python -m tests.create_sample_template
  python tests/create_sample_template.py
"""
import sys
from pathlib import Path

if __name__ == "__main__" and __package__ is None:
    _root = Path(__file__).resolve().parent.parent
    sys.path.insert(0, str(_root))

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_contract_template(output_path: Path) -> None:
    """Шаблон «Договор оказания услуг» — поля: CLIENT_NAME, CONTRACT_NUMBER, CONTRACT_DATE, CONTRACT_TYPE, AMOUNT."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(12)
    style.font.name = "Times New Roman"

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ДОГОВОР ОКАЗАНИЯ УСЛУГ")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("№ ")
    p.add_run("{{ CONTRACT_NUMBER }}")
    p.add_run(" от ")
    p.add_run("{{ CONTRACT_DATE }}")
    p.add_run(" г.")

    doc.add_paragraph()
    doc.add_paragraph(
        "Настоящий договор заключён между Исполнителем и Заказчиком в лице "
        "{{ CLIENT_NAME }}, именуемым в дальнейшем «Заказчик»."
    )
    doc.add_paragraph()
    doc.add_paragraph("Тип договора: {{ CONTRACT_TYPE }}.")
    doc.add_paragraph("Сумма договора: {{ AMOUNT }} руб.")
    doc.add_paragraph()
    doc.add_paragraph("Подписи сторон: _______________________  _______________________")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Создан шаблон: {output_path}")


def create_act_template(output_path: Path) -> None:
    """Шаблон «Акт выполненных работ» — поля: ACT_NUMBER, ACT_DATE, CLIENT_NAME, EXECUTOR_NAME, WORK_DESCRIPTION."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(12)
    style.font.name = "Times New Roman"

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("АКТ ВЫПОЛНЕННЫХ РАБОТ")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("№ ")
    p.add_run("{{ ACT_NUMBER }}")
    p.add_run(" от ")
    p.add_run("{{ ACT_DATE }}")

    doc.add_paragraph()
    doc.add_paragraph("Заказчик: {{ CLIENT_NAME }}")
    doc.add_paragraph("Исполнитель: {{ EXECUTOR_NAME }}")
    doc.add_paragraph()
    doc.add_paragraph("Описание выполненных работ: {{ WORK_DESCRIPTION }}")
    doc.add_paragraph()
    doc.add_paragraph("Подписи: _______________________  _______________________")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Создан шаблон: {output_path}")


def create_invoice_loop_template(output_path: Path) -> None:
    """
    Шаблон «Счёт-фактура с позициями» — простые поля + цикл items.
    Простые поля: INVOICE_NUMBER, INVOICE_DATE, CLIENT_NAME.
    Цикл: {% for item in items %} — item.NAME, item.QTY, item.PRICE.
    """
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(12)
    style.font.name = "Times New Roman"

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("СЧЁТ-ФАКТУРА")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph("№ {{ INVOICE_NUMBER }} от {{ INVOICE_DATE }}")
    doc.add_paragraph("Покупатель: {{ CLIENT_NAME }}")
    doc.add_paragraph()
    doc.add_paragraph("Перечень работ / услуг:")
    # docxtpl: весь цикл должен быть в одном параграфе (одном run)
    doc.add_paragraph(
        "{% for item in items %}\n{{ item.NAME }}  |  {{ item.QTY }} шт.  |  {{ item.PRICE }} руб.\n{% endfor %}"
    )
    doc.add_paragraph()
    doc.add_paragraph("Исполнитель: _______________________")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Создан шаблон: {output_path}")


def create_mixed_loop_template(output_path: Path) -> None:
    """
    Шаблон «Акт сдачи-приёмки с таблицей работ» — простые поля + цикл works.
    Простые поля: EXECUTOR, CONTRACT_DATE.
    Цикл: {% for row in works %} — row.WORK, row.AMOUNT.
    """
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(12)
    style.font.name = "Times New Roman"

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("АКТ СДАЧИ-ПРИЁМКИ РАБОТ")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph("Исполнитель: {{ EXECUTOR }}")
    doc.add_paragraph("Дата: {{ CONTRACT_DATE }}")
    doc.add_paragraph()
    doc.add_paragraph("Состав выполненных работ:")
    # docxtpl: весь цикл в одном параграфе
    doc.add_paragraph(
        "{% for row in works %}\n{{ row.WORK }}  —  {{ row.AMOUNT }} руб.\n{% endfor %}"
    )
    doc.add_paragraph()
    doc.add_paragraph("Подписи сторон: _______________________  _______________________")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Создан шаблон: {output_path}")


def create_invoice_simple_template(output_path: Path) -> None:
    """Шаблон «Счёт на оплату» — только простые поля: INVOICE_NO, DATE, CUSTOMER, TOTAL."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(12)
    style.font.name = "Times New Roman"

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("СЧЁТ НА ОПЛАТУ")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph("№ {{ INVOICE_NO }} от {{ DATE }}")
    doc.add_paragraph()
    doc.add_paragraph("Заказчик: {{ CUSTOMER }}")
    doc.add_paragraph()
    doc.add_paragraph("Сумма к оплате: {{ TOTAL }} руб.")
    doc.add_paragraph()
    doc.add_paragraph("Подпись: _______________________")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Создан шаблон: {output_path}")


def create_delivery_note_template(output_path: Path) -> None:
    """
    Шаблон «Накладная» — простые поля + цикл goods (таблица товаров).
    Простые поля: CONSIGNEE, DELIVERY_DATE.
    Цикл: {% for item in goods %} — item.NAME, item.QTY, item.UNIT.
    """
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(12)
    style.font.name = "Times New Roman"

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("НАКЛАДНАЯ")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph("Получатель: {{ CONSIGNEE }}")
    doc.add_paragraph("Дата отгрузки: {{ DELIVERY_DATE }}")
    doc.add_paragraph()
    doc.add_paragraph("Перечень товаров:")
    # docxtpl: весь цикл в одном параграфе
    doc.add_paragraph(
        "{% for item in goods %}\n{{ item.NAME }}  |  {{ item.QTY }}  {{ item.UNIT }}\n{% endfor %}"
    )
    doc.add_paragraph()
    doc.add_paragraph("Сдал: _______________________   Принял: _______________________")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Создан шаблон: {output_path}")


def create_project_report_template(output_path: Path) -> None:
    """
    Шаблон «Отчёт по проекту» — 3 различных таблицы/цикла в одном документе.
    Простые поля: PROJECT_NAME, REPORT_DATE, AUTHOR.
    Цикл 1: participants (item) — NAME, ROLE, CONTACT.
    Цикл 2: milestones (m) — MILESTONE, DEADLINE, STATUS.
    Цикл 3: budget_lines (line) — DESCRIPTION, AMOUNT, CATEGORY.
    """
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(12)
    style.font.name = "Times New Roman"

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ОТЧЁТ ПО ПРОЕКТУ")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph("Проект: {{ PROJECT_NAME }}")
    doc.add_paragraph("Дата отчёта: {{ REPORT_DATE }}")
    doc.add_paragraph("Автор: {{ AUTHOR }}")
    doc.add_paragraph()

    # docxtpl: каждый цикл — в одном параграфе (нельзя разбивать for/endfor по параграфам)
    doc.add_paragraph("1. Участники проекта:")
    doc.add_paragraph(
        "{% for item in participants %}\n{{ item.NAME }}  —  {{ item.ROLE }}  ({{ item.CONTACT }})\n{% endfor %}"
    )
    doc.add_paragraph()

    doc.add_paragraph("2. Этапы и сроки:")
    doc.add_paragraph(
        "{% for m in milestones %}\n{{ m.MILESTONE }}  |  Срок: {{ m.DEADLINE }}  |  Статус: {{ m.STATUS }}\n{% endfor %}"
    )
    doc.add_paragraph()

    doc.add_paragraph("3. Бюджет:")
    doc.add_paragraph(
        "{% for line in budget_lines %}\n{{ line.DESCRIPTION }}  —  {{ line.AMOUNT }} руб.  [{{ line.CATEGORY }}]\n{% endfor %}"
    )
    doc.add_paragraph()
    doc.add_paragraph("Подпись: _______________________")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Создан шаблон: {output_path}")


def main():
    root = Path(__file__).resolve().parent
    out_dir = root / "fixtures"
    # Простые шаблоны (только поля)
    create_contract_template(out_dir / "contract_services.docx")
    create_act_template(out_dir / "act_works.docx")
    create_invoice_simple_template(out_dir / "invoice_simple.docx")
    # Шаблоны с циклами / таблицами
    create_invoice_loop_template(out_dir / "invoice_loop.docx")
    create_mixed_loop_template(out_dir / "mixed_loop.docx")
    create_delivery_note_template(out_dir / "delivery_note.docx")
    create_project_report_template(out_dir / "project_report.docx")
    print("Готово. Файлы в каталоге tests/fixtures/")
    print("В приложении: «Загрузить шаблон...» -> указать один из этих файлов.")


if __name__ == "__main__":
    main()
